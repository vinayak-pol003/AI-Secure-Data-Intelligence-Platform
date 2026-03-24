import io
from typing import Dict


def parse_doc(content_bytes: bytes, filename: str = "", ext: str = "docx") -> Dict:
    """
    Parse a .doc or .docx file and extract text content.
    Returns structured content with paragraph and table data.
    """
    content = ""
    paragraph_count = 0
    table_count = 0

    if ext == "docx":
        try:
            from docx import Document
            doc = Document(io.BytesIO(content_bytes))

            sections = []

            # Extract paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            paragraph_count = len(paragraphs)
            if paragraphs:
                sections.append("\n".join(paragraphs))

            # Extract tables
            for i, table in enumerate(doc.tables, 1):
                table_count += 1
                table_lines = [f"--- TABLE {i} ---"]
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        table_lines.append(row_text)
                if len(table_lines) > 1:
                    sections.append("\n".join(table_lines))

            content = "\n\n".join(sections)

        except ImportError:
            content = content_bytes.decode("utf-8", errors="ignore")

        except Exception:
            content = content_bytes.decode("utf-8", errors="ignore")

    elif ext == "doc":
        # .doc (old Word format) - try antiword or textract
        try:
            import subprocess
            result = subprocess.run(
                ["antiword", "-"],
                input=content_bytes,
                capture_output=True,
                timeout=10
            )
            content = result.stdout.decode("utf-8", errors="ignore")
        except Exception:
            # Fallback: raw decode (may have noise for binary .doc)
            content = content_bytes.decode("utf-8", errors="ignore")

    else:
        content = content_bytes.decode("utf-8", errors="ignore")

    lines = content.splitlines()

    return {
        "content": content,
        "lines": lines,
        "line_count": len(lines),
        "paragraph_count": paragraph_count,
        "table_count": table_count,
        "filename": filename,
        "file_type": ext
    }