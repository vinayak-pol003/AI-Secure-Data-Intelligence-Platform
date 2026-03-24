from docx import Document
import json

doc = Document(r"C:\Users\LENOVO\OneDrive\Desktop\SISA Project\AI Secure Data Intelligence Platform.docx")

output = []
for para in doc.paragraphs:
    if para.text.strip():
        output.append({"style": para.style.name, "text": para.text})

for i, table in enumerate(doc.tables):
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        output.append({"style": f"table_{i}", "text": " | ".join(cells)})

print(json.dumps(output, ensure_ascii=False, indent=2))
